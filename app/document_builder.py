import re
import base64
import logging
from io import BytesIO
from datetime import datetime

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

"""
Ported from the Streamlit app's document generation logic (create_a4_html,
create_word_docx, html_to_pdf) so the FastAPI backend produces identical
output — including the Hindi/Devanagari font fix (Noto Sans Devanagari +
bidi language tag) that took a few rounds to get right there.
"""

class StoredLogo(BytesIO):
    def __init__(self, data: bytes, mimetype: str):
        super().__init__(data)
        self.type = mimetype

def clean_math_for_word(text):
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\\((.*?)\\\)', r'\1', text)
    text = re.sub(r'\\\[(.*?)\\\]', r'\1', text)
    latex_map = {r'\pi': 'π', r'\theta': 'θ', r'\sqrt': '√', r'\times': '×', r'\div': '÷', '$': '', '^2': '²', '^3': '³'}
    for k, v in latex_map.items(): text = text.replace(k, v)
    text = text.replace('☐', '[ ]').replace('☑', '[x]').replace('•', '-').replace('◦', '-')
    text = text.replace('\u200b', '').replace('\u2022', '-').replace('\u25cf', '-').replace('\u25cb', '-')
    return text.strip()

# 🌟 HTML RENDERER 🌟

def create_a4_html(md_content, i_name, i_address, i_contact, t_name, inst_logo=None, is_2_col=False, sub="Subject", grade="Class", total_m="Marks", exam_time="Time", topics="", custom_instructions="", reading_time=""):
    md_content = clean_math_for_word(md_content)
    
    md_content = re.sub(r"^#.*?\*\*\*", "", md_content, count=1, flags=re.DOTALL).strip()
    md_content = re.sub(r"^\*\*Subject:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Class:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Marks:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Time:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    
    md_content = re.sub(r"^\d+\.\s", "**Q.** ", md_content, flags=re.MULTILINE)
    md_content = md_content.strip()
    
    logo_html_inline = ""
    logo_footer = ""
    if inst_logo:
        inst_logo.seek(0)
        b64 = base64.b64encode(inst_logo.getvalue()).decode()
        logo_html_inline = f"<td style='width: 65px; padding-right: 15px; vertical-align: middle;'><img src='data:{inst_logo.type};base64,{b64}' style='max-height: 55px; max-width: 65px;'/></td>"
        logo_footer = f"<img src='data:{inst_logo.type};base64,{b64}' style='height: 18px; vertical-align: middle; margin-right: 8px;'/>"
    
    main_heading_text = topics.strip().upper() if topics.strip() != "" else sub.upper()
    
    reading_time_html = f" (+ {reading_time} reading time)" if reading_time.strip() else ""
    custom_instructions_html = f"""
    <div style='border: 1px solid #999; padding: 8px 12px; margin-bottom: 15px; font-size: 13px; background: #fafafa;'>
        <strong>Instructions:</strong> {custom_instructions.strip()}
    </div>
    """ if custom_instructions.strip() else ""

    custom_header = f"""
    <div style='border-bottom: 2px solid black; padding-bottom: 10px; margin-bottom: 10px; width: 100%;'>
        <table style='width: 100%; border-collapse: collapse; border: none; margin-bottom: 10px;'>
            <tr>
                <td style='text-align: center; vertical-align: middle; border: none;'>
                    <table style='margin: 0 auto;'>
                        <tr>
                            {logo_html_inline}
                            <td style='vertical-align: middle;'>
                                <h1 style='margin: 0; font-size: 24px; font-family: "Noto Sans", "Nirmala UI", "Times New Roman", serif; font-weight: 900; text-transform: uppercase; white-space: nowrap;'>{i_name}</h1>
                            </td>
                        </tr>
                    </table>
                </td>
            </tr>
        </table>
        <table style='width: 100%; font-weight: bold; font-size: 13px; border: none;'>
            <tr>
                <td style='text-align: left; vertical-align: bottom; width: 33%; border: none;'>Class : {grade}<br>Time : {exam_time}{reading_time_html}</td>
                <td style='text-align: center; vertical-align: middle; width: 34%; border: none;'>
                    <div style='border: 2px solid black; border-radius: 12px; display: inline-block; padding: 4px 25px; font-weight: bold; font-size: 14px; background: white;'>
                        EXAMINATION
                    </div>
                </td>
                <td style='text-align: right; vertical-align: bottom; width: 33%; border: none;'>Sub.: {sub}<br>Marks: {total_m}</td>
            </tr>
        </table>
    </div>
    <div style='border-top: 1px solid black; border-bottom: 3px solid black; padding: 2px 0; margin-bottom: 15px;'>
        <div style='background-color: black; color: white; padding: 5px; text-align: center; font-weight: bold; font-size: 15px; text-transform: uppercase; letter-spacing: 1px;'>
            Multiple Choice Questions & Theory
        </div>
    </div>
    <h2 style='text-align: center; text-decoration: underline; text-transform: uppercase; margin-top: 0; margin-bottom: 15px; font-size: 18px;'>{main_heading_text}</h2>
    {custom_instructions_html}
    """

    # A SIMPLER header for the Answer Key page — reusing the exact same
    # complex nested-table header twice in one document (verbatim) triggers
    # a layout/text-corruption bug in xhtml2pdf's table engine (confirmed
    # independent of language/content — happens even with plain English).
    simple_ak_header = f"""
    <div style='border-bottom: 2px solid black; padding-bottom: 8px; margin-bottom: 15px; text-align: center;'>
        <span style='font-size: 20px; font-weight: 900; text-transform: uppercase;'>{i_name}</span>
    </div>
    """

    footer_html = f"""
    <div class="footer-content print-only-footer">
        {logo_footer}<strong>{i_name}</strong> | 📍 {i_address} | 📞 {i_contact} | 👨‍🏫 <strong>{t_name}</strong>
    </div>
    """

    # This watermark+footer pair lives OUTSIDE the table (as siblings within
    # a .page-section wrapper), used for BOTH on-screen viewing and print —
    # position:fixed CSS Paged Media repetition turned out unreliable across
    # real browsers' Print/Ctrl+P engines (only appeared once, near the end,
    # in testing), even though it worked correctly in our own WeasyPrint
    # PDF export. This per-section approach doesn't depend on that browser
    # behavior at all. Each "page" (main content, and the Answer Key if
    # present) gets its own copy, so both appear at the end of EVERY
    # section, not just once at the very end. The footer still also has a
    # separate print-only copy inside the table's tfoot, since that one
    # continues to repeat correctly per physical page during printing.
    def screen_overlay_html():
        return f"""
        <div class="screen-watermark">{i_name}</div>
        <div class="footer-content screen-only-footer">
            {logo_footer}<strong>{i_name}</strong> | 📍 {i_address} | 📞 {i_contact} | 👨‍🏫 <strong>{t_name}</strong>
        </div>
        """

    ans_split_marker = "|||ANSWER_KEY_SPLIT|||"
    md_content = re.sub(r'(?im)^#+\s*Answer Key.*$', ans_split_marker, md_content)

    if ans_split_marker in md_content:
        q_part, a_part = md_content.split(ans_split_marker)
        # Two SEPARATE tables (not one table with a page-break inserted mid-row) —
        # putting a manual page-break inside a single table's body row corrupts
        # text rendering in xhtml2pdf (confirmed independent of language/content).
        # Splitting into two self-contained tables, with page-break-before on
        # the second table itself, avoids that bug entirely.
        tables_html = f"""
        <div class="page-section">
            <table>
                <thead><tr><td></td></tr></thead>
                <tbody><tr><td>
                    {custom_header}
                    <div class="content-body">{markdown.markdown(q_part.strip())}</div>
                </td></tr></tbody>
                <tfoot><tr><td>{footer_html}</td></tr></tfoot>
            </table>
            {screen_overlay_html()}
        </div>
        <div class="page-section">
            <table style="page-break-before: always;">
                <thead><tr><td></td></tr></thead>
                <tbody><tr><td>
                    {simple_ak_header}
                    <h2 style="text-align: center; text-decoration: underline; margin-bottom: 15px;">ANSWER KEY</h2>
                    <div class="content-body">{markdown.markdown(a_part.strip())}</div>
                </td></tr></tbody>
                <tfoot><tr><td>{footer_html}</td></tr></tfoot>
            </table>
            {screen_overlay_html()}
        </div>
        """
    else:
        tables_html = f"""
        <div class="page-section">
            <table>
                <thead><tr><td></td></tr></thead>
                <tbody><tr><td>
                    {custom_header}
                    <div class="content-body">{markdown.markdown(md_content.strip())}</div>
                </td></tr></tbody>
                <tfoot><tr><td>{footer_html}</td></tr></tfoot>
            </table>
            {screen_overlay_html()}
        </div>
        """

    col_style = "column-count: 2; column-gap: 15mm; column-rule: 1px solid #000; font-size: 14px;" if is_2_col else "font-size: 16px;"

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
    body {{ background: #f0f0f0; font-family: 'Noto Sans', 'Nirmala UI', 'Times New Roman', serif; margin: 0; padding: 20px; display: flex; justify-content: center; }} 
    .a4-page {{ background: white; width: 210mm; min-height: 297mm; padding: 20px; box-shadow: 0 0 10px rgba(0,0,0,0.2); box-sizing: border-box; position: relative; overflow: hidden; }} 
    table {{ width: 100%; border-collapse: collapse; border: none; position: relative; z-index: 1; }}
    td {{ border: none; padding: 0; }}
    h1, h2, h3 {{ text-align: center; column-span: all; }} 
    h2 {{ font-size: 16px; border-bottom: 1px dashed #ccc; padding-bottom: 5px; }}
    .content-body {{ {col_style} position: relative; z-index: 1; text-align: justify; }} 
    .content-body p {{ margin-bottom: 8px; margin-top: 4px; }}
    /* Each "page" (.page-section) gets its own watermark + footer, used
       for BOTH on-screen viewing and print/PDF — a single approach that
       works reliably everywhere, rather than depending on position:fixed
       CSS Paged Media repetition, which behaves inconsistently: WeasyPrint
       (our own "Download PDF" button) repeats it correctly per page, but
       real browsers' own Print/Ctrl+P engines were confirmed NOT to repeat
       it reliably (it only showed up once, near the end of the document).
       position:relative on .page-section gives its absolutely positioned
       children (watermark, footer) a sane containing block and a real
       stacking context, so z-index:0 vs z-index:1 below resolve
       predictably regardless of viewing context. */
    .page-section {{ position: relative; min-height: 297mm; display: flex; flex-direction: column; }} .page-section > table {{ flex: 1 0 auto; }}
    .screen-watermark {{ position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%) rotate(-45deg); font-size: 85px; color: rgba(0, 0, 0, 0.06); z-index: 0; pointer-events: none; white-space: nowrap; font-weight: bold; text-transform: uppercase; }}
    .footer-content {{ text-align: center; padding-top: 10px; border-top: 2px dashed #bbb; font-size: 13px; color: #444; z-index: 1; background: white; }}
    .screen-only-footer {{ }}
    .print-only-footer {{ display: none; }}
    @media print {{ 
        @page {{ size: A4; margin: 0; }} 
        body {{ background: white; padding: 0; margin: 0; display: block; }} 
        .a4-page {{ box-shadow: none; width: 100%; min-height: auto; padding: 10mm; margin: 0; page-break-after: always; }} 
        .screen-watermark {{ color: rgba(0, 0, 0, 0.06) !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
        tfoot {{ display: table-footer-group; }}
        .screen-only-footer {{ display: none; }}
        .print-only-footer {{ display: block; margin-top: 20px; }}
    }} 
    </style></head><body><div class="a4-page">
    {tables_html}
    </div></body></html>"""

import os
import threading

_FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")
# Serializes PDF generation. WeasyPrint's font loading/subsetting isn't
# guaranteed thread-safe, and on a memory-constrained server, concurrent
# PDF requests could plausibly interfere with each other's font embedding
# (some Hindi glyphs turning into tofu boxes only in production, never
# reproducible locally, pointed at exactly this kind of shared-state issue).
_PDF_LOCK = threading.Lock()

def html_to_pdf(html_string):
    try:
        # Imported lazily (not at module load time) so that a missing
        # system dependency (GTK3, needed by WeasyPrint on Windows) only
        # breaks PDF export specifically — not the entire backend's ability
        # to start up, which would otherwise take down login, paper
        # generation, Word/HTML export, everything, just because PDF
        # generation's native library isn't installed on this machine.
        from weasyprint import HTML
        # Swap emoji for text labels in the footer — our bundled font
        # doesn't include emoji glyphs (HTML/Word keep the emoji, since
        # browsers/Word render them fine using their own system fonts).
        html_string = (
            html_string
            .replace("📍", "Address:")
            .replace("📞", "Phone:")
            .replace("👨‍🏫", "Teacher:")
        )
        # PDF generation happens on the server, which has no access to the
        # reader's installed fonts — any font used must be embedded in the
        # PDF itself. We bundle a merged font (Noto Sans + Noto Sans
        # Devanagari) covering both English and Hindi glyphs, so this works
        # reliably regardless of what's installed on the deployment server.
        regular_latin = os.path.join(_FONTS_DIR, "NotoSans-Regular.ttf")
        regular_dev = os.path.join(_FONTS_DIR, "NotoSansDevanagari-Regular.ttf")
        bold_latin = os.path.join(_FONTS_DIR, "NotoSans-Bold.ttf")
        bold_dev = os.path.join(_FONTS_DIR, "NotoSansDevanagari-Bold.ttf")
        # Two font files per weight (Latin + Devanagari), selected
        # automatically per-character via unicode-range, instead of one
        # merged font file. The earlier merged-font approach (built via
        # fonttools pyftmerge) is a plausible source of the intermittent
        # Hindi-glyph corruption seen only in production and never
        # reproducible locally — merging two fonts' cmap/glyph tables is a
        # non-trivial transformation. Using the original, unmodified Google
        # Noto font files directly removes that step entirely.
        pdf_font_override = f"""
        <style>
            @font-face {{ font-family: 'PaperBanaoPDF'; src: url({regular_latin}); font-weight: normal; unicode-range: U+0000-08FF, U+2000-23FF; }}
            @font-face {{ font-family: 'PaperBanaoPDF'; src: url({regular_dev}); font-weight: normal; unicode-range: U+0900-097F; }}
            @font-face {{ font-family: 'PaperBanaoPDF'; src: url({bold_latin}); font-weight: bold; unicode-range: U+0000-08FF, U+2000-23FF; }}
            @font-face {{ font-family: 'PaperBanaoPDF'; src: url({bold_dev}); font-weight: bold; unicode-range: U+0900-097F; }}
            * {{ font-family: 'PaperBanaoPDF' !important; }}
        </style>
        """
        html_for_pdf = html_string.replace("</head>", pdf_font_override + "</head>")
        # WeasyPrint (built on Pango) correctly shapes complex Indic scripts
        # — conjunct consonants, matras attaching to the right glyph, etc.
        # xhtml2pdf's reportlab engine has NO text-shaping support at all,
        # which is why Hindi text there showed as tofu/missing-glyph boxes
        # for any word using a conjunct — this isn't a font problem, it's a
        # rendering-engine capability gap that no font file can work around.
        with _PDF_LOCK:
            # full_fonts=True embeds the complete font file instead of a
            # subset containing only "used" characters. WeasyPrint's default
            # subsetting has a bug with complex Devanagari conjunct glyphs
            # (formed via ligature substitution, not a simple 1:1 character
            # mapping) — the subsetter sometimes drops glyphs those
            # conjuncts need, so the PDF's text layer is correct Unicode
            # but the visual glyph is missing (tofu box) since the embedded
            # font subset doesn't actually contain it. This trades a
            # somewhat larger PDF file for correct rendering.
            pdf_bytes = HTML(string=html_for_pdf).write_pdf(full_fonts=True)
        return pdf_bytes
    except Exception as e:
        logging.error(f"[PDF Generation Error] {e}")
        return None

# 🌟 WORD RENDERER 🌟
def create_word_docx(md_content, i_name, i_address, i_contact, t_name, inst_logo=None, is_2_col=False, sub="Subject", grade="Class", total_m="Marks", exam_time="Time", topics="", custom_instructions="", reading_time=""):
    doc = Document()
    
    md_content = re.sub(r"^#.*?\*\*\*", "", md_content, count=1, flags=re.DOTALL).strip()
    md_content = re.sub(r"^\*\*Subject:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Class:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Marks:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Time:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\d+\.\s", "**Q.** ", md_content, flags=re.MULTILINE)
    md_content = md_content.strip()
        
    md_content = md_content.replace('\r', '')
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial' 
    font.size = Pt(11)
    
    rFonts = style.element.rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn('w:cs'), 'Noto Sans Devanagari') 
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
    style_lang = style.element.rPr.find(qn('w:lang'))
    if style_lang is None:
        style_lang = style.element.rPr.makeelement(qn('w:lang'), {})
        style.element.rPr.append(style_lang)
    style_lang.set(qn('w:bidi'), 'hi-IN')
    
    for i in range(3):
        try:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Arial'
            if h_style.element.rPr.rFonts is not None:
                h_style.element.rPr.rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
                h_style.element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                h_style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
            # Same fix as the Normal style above: without this explicit lang
            # tag directly on the heading style, Word doesn't reliably route
            # Hindi/Devanagari text in headings through the cs font, even
            # though the font name above is set correctly — showing tofu
            # boxes for Hindi section headers specifically.
            h_style_lang = h_style.element.rPr.find(qn('w:lang'))
            if h_style_lang is None:
                h_style_lang = h_style.element.rPr.makeelement(qn('w:lang'), {})
                h_style.element.rPr.append(h_style_lang)
            h_style_lang.set(qn('w:bidi'), 'hi-IN')
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 0:
                h_style.font.size = Pt(16)
                h_style.font.bold = True
            elif i == 1:
                h_style.font.size = Pt(12)
                h_style.font.bold = True
            elif i == 2:
                h_style.font.size = Pt(11)
                h_style.font.bold = True
        except KeyError: pass

    if is_2_col:
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.4)

    def apply_cs_font(run):
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(qn('w:cs'), 'Noto Sans Devanagari')
        rfonts.set(qn('w:ascii'), 'Arial')
        rfonts.set(qn('w:hAnsi'), 'Arial')
        # Without an explicit language tag, Word doesn't reliably classify
        # Devanagari text as "complex script" and may render it with the
        # ascii font (Arial, no Devanagari glyphs = tofu boxes) regardless
        # of the w:cs font specified above. This tag is what makes Word
        # actually route the text correctly.
        lang = rpr.find(qn('w:lang'))
        if lang is None:
            lang = rpr.makeelement(qn('w:lang'), {})
            rpr.append(lang)
        lang.set(qn('w:bidi'), 'hi-IN')

    def insert_chate_header():
        title_table = doc.add_table(rows=1, cols=1)
        p1 = title_table.cell(0,0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if inst_logo is not None:
            try:
                inst_logo.seek(0)
                r_logo = p1.add_run()
                r_logo.add_picture(inst_logo, height=Inches(0.38))
                p1.add_run("   ") 
            except Exception: pass
            
        r1 = p1.add_run(i_name.upper())
        r1.bold = True
        r1.font.size = Pt(18)
        apply_cs_font(r1)
        
        details_table = doc.add_table(rows=1, cols=3)
        details_table.autofit = False
        for cell in details_table.columns[0].cells: cell.width = Inches(2.0)
        for cell in details_table.columns[1].cells: cell.width = Inches(3.0)
        for cell in details_table.columns[2].cells: cell.width = Inches(2.0)

        p3 = details_table.cell(0,0).paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        time_text = f"Time : {exam_time}" + (f" (+ {reading_time} reading time)" if reading_time.strip() else "")
        r3 = p3.add_run(f"Class : {grade}\n{time_text}")
        r3.bold = True
        r3.font.size = Pt(10)
        apply_cs_font(r3)

        p4 = details_table.cell(0,1).paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run("\n[ EXAMINATION ]")
        r4.bold = True
        r4.font.size = Pt(12)
        apply_cs_font(r4)

        p2 = details_table.cell(0,2).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(f"Sub.: {sub}\nMarks: {total_m}")
        r2.bold = True
        r2.font.size = Pt(10)
        apply_cs_font(r2)
        
        doc.add_paragraph("__________________________________________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt = doc.add_paragraph("MULTIPLE CHOICE QUESTIONS & THEORY")
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.runs[0].bold = True
        apply_cs_font(pt.runs[0])
        
        main_heading_text = topics.strip().upper() if topics.strip() != "" else sub.upper()
        ptopics = doc.add_paragraph(main_heading_text)
        ptopics.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ptopics.runs[0].underline = True
        ptopics.runs[0].font.size = Pt(14)
        ptopics.runs[0].bold = True
        apply_cs_font(ptopics.runs[0])

        if custom_instructions.strip():
            pinst = doc.add_paragraph()
            rinst_label = pinst.add_run("Instructions: ")
            rinst_label.bold = True
            rinst_label.font.size = Pt(10)
            apply_cs_font(rinst_label)
            rinst = pinst.add_run(custom_instructions.strip())
            rinst.font.size = Pt(10)
            apply_cs_font(rinst)

        doc.add_paragraph() 

    insert_chate_header()

    if is_2_col:
        new_section = doc.add_section(0) 
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720') 

    for line in md_content.split('\n'):
        line_clean = line.strip()
        if not line_clean: continue
        line_clean = clean_math_for_word(line_clean)
        
        if "Answer Key" in line_clean or "ANSWER KEY" in line_clean:
            doc.add_page_break() 
            insert_chate_header() 
            ak_heading = doc.add_heading("Answer Key", level=1)
            for r in ak_heading.runs:
                apply_cs_font(r)
            continue
            
        if line_clean.startswith('# '): 
            h = doc.add_heading(line_clean.replace('# ', ''), level=1)
            for r in h.runs:
                apply_cs_font(r)
        elif line_clean.startswith('## '): 
            h = doc.add_heading(line_clean.replace('## ', ''), level=2)
            for r in h.runs:
                apply_cs_font(r)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            parts = re.split(r'\*\*(.*?)\*\*', line_clean)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: run.bold = True
                apply_cs_font(run)
                
    if doc.sections:
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if inst_logo is not None:
            try:
                inst_logo.seek(0)
                run_logo = footer_para.add_run()
                run_logo.add_picture(inst_logo, height=Inches(0.18))
                footer_para.add_run("  ") 
            except Exception: pass
            
        run_name = footer_para.add_run(f"{i_name}  |  ")
        run_name.font.size = Pt(10)
        run_name.font.bold = True
        run_name.font.color.rgb = RGBColor(100, 100, 100)
        apply_cs_font(run_name)
        
        run_rest = footer_para.add_run(f"📍 {i_address}  |  📞 {i_contact}  |  👨‍🏫 {t_name}")
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(100, 100, 100)
        apply_cs_font(run_rest)
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
